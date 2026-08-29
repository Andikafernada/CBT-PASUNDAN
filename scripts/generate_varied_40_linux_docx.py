import os
import sys
import shutil

if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageDraw

TEMP_DIR = os.path.join(os.path.dirname(__file__), "temp_varied_images")
os.makedirs(TEMP_DIR, exist_ok=True)

def create_terminal_image(filename, q_type, category, title, terminal_lines, bg_color="#0f172a", border_color="#38bdf8"):
    width, height = 750, 240
    img = Image.new("RGB", (width, height), color=bg_color)
    draw = ImageDraw.Draw(img)

    # Outer border
    draw.rectangle([8, 8, width - 8, height - 8], outline=border_color, width=2)

    # Terminal Top Bar
    draw.rectangle([8, 8, width - 8, 38], fill="#1e293b")
    # Red, Yellow, Green window buttons
    draw.ellipse([20, 18, 30, 28], fill="#ef4444")
    draw.ellipse([36, 18, 46, 28], fill="#f59e0b")
    draw.ellipse([52, 18, 62, 28], fill="#10b981")
    draw.text((80, 15), f"root@debian12:~# [{q_type}] {category} - {title}", fill="#94a3b8")

    # Terminal Code Lines
    y = 52
    for line in terminal_lines:
        if line.startswith("root@") or line.startswith("$") or line.startswith("#"):
            draw.text((25, y), line, fill="#38bdf8")
        elif "OK" in line or "active (running)" in line or "SUCCESS" in line or "Benar" in line:
            draw.text((25, y), line, fill="#4ade80")
        elif "ERROR" in line or "FAILED" in line or "FAIL" in line or "Salah" in line:
            draw.text((25, y), line, fill="#f87171")
        elif line.startswith(";") or line.startswith("//"):
            draw.text((25, y), line, fill="#64748b")
        else:
            draw.text((25, y), line, fill="#f8fafc")
        y += 26

    path = os.path.join(TEMP_DIR, filename)
    img.save(path, format="PNG")
    return path

# 40 Questions across MULTIPLE VARIETIES:
# 1-15: MULTIPLE_CHOICE (PG)
# 16-25: COMPLEX_MULTIPLE_CHOICE (PG Kompleks / Checklist)
# 26-30: TRUE_FALSE (Benar / Salah)
# 31-35: MATCHING (Menjodohkan)
# 36-40: ESSAY (Uraian / Studi Kasus)

QUESTIONS = [
    # =========================================================================
    # BAGIAN 1: PILIHAN GANDA BIASA (Soal 1 - 15)
    # =========================================================================
    {
        "type": "PG",
        "cat": "VIRTUALBOX",
        "title": "Mode Jaringan Host-Only Adapter",
        "q": "Pada VirtualBox, administrator ingin agar VM Debian 12 dapat berkomunikasi dua arah dengan OS Host (Windows) tanpa terhubung ke internet luar, serta memiliki subnet IP statis tersendiri. Mode adapter jaringan VirtualBox yang paling tepat digunakan adalah...",
        "lines": [
            "Settings > Network > Adapter 1",
            "Attached to: [ Host-only Adapter ]",
            "Name: VirtualBox Host-Only Ethernet Adapter",
            "Promiscuous Mode: Allow All"
        ],
        "options": [
            "A. Host-only Adapter",
            "B. NAT (Network Address Translation)",
            "C. Bridged Adapter",
            "D. Internal Network",
            "E. Generic Driver"
        ],
        "kunci": "A"
    },
    {
        "type": "PG",
        "cat": "VIRTUALBOX",
        "title": "Port Forwarding NAT Apache",
        "q": "Jika VM Debian 12 menggunakan adapter jaringan NAT pada VirtualBox dan menjalankan web server Apache di port 80, konfigurasi Port Forwarding agar web server dapat diakses dari browser Host melalui alamat http://localhost:8080 adalah...",
        "lines": [
            "Rule 1: Name: HTTP-Web, Protocol: TCP",
            "Host IP: 127.0.0.1, Host Port: 8080",
            "Guest IP: 10.0.2.15, Guest Port: 80",
            "Status: Active Forwarding"
        ],
        "options": [
            "A. Host Port: 8080, Guest Port: 80",
            "B. Host Port: 80, Guest Port: 8080",
            "C. Host Port: 22, Guest Port: 80",
            "D. Host Port: 443, Guest Port: 8080",
            "E. Host Port: 53, Guest Port: 53"
        ],
        "kunci": "A"
    },
    {
        "type": "PG",
        "cat": "DEBIAN 12 CLI",
        "title": "Konfigurasi IP Statis (/etc/network/interfaces)",
        "q": "Perhatikan konfigurasi file /etc/network/interfaces pada Debian 12 berikut:\nPerintah yang tepat untuk menerapkan perubahan konfigurasi IP address tersebut tanpa melakukan restart server secara keseluruhan adalah...",
        "lines": [
            "auto ens18",
            "iface ens18 inet static",
            "    address 192.168.10.1/24",
            "    gateway 192.168.10.254",
            "    dns-nameservers 192.168.10.1 8.8.8.8"
        ],
        "options": [
            "A. systemctl restart networking",
            "B. systemctl restart apache2",
            "C. ifconfig ens18 restart",
            "D. ip address reload ens18",
            "E. service network-manager stop"
        ],
        "kunci": "A"
    },
    {
        "type": "PG",
        "cat": "LINUX PERMISSIONS",
        "title": "Hak Akses File (chmod 755 & chown)",
        "q": "Administrator ingin memberikan hak akses penuh (baca, tulis, eksekusi) untuk pemilik (owner), serta hanya hak baca dan eksekusi untuk grup dan user lainnya pada folder /var/www/html. Nilai numerik chmod yang tepat adalah...",
        "lines": [
            "root@debian12:~# chmod 755 -R /var/www/html",
            "root@debian12:~# chown -R www-data:www-data /var/www/html",
            "drwxr-xr-x 2 www-data www-data 4096 Aug 29 07:00 /var/www/html"
        ],
        "options": [
            "A. chmod 755 /var/www/html",
            "B. chmod 777 /var/www/html",
            "C. chmod 644 /var/www/html",
            "D. chmod 700 /var/www/html",
            "E. chmod 600 /var/www/html"
        ],
        "kunci": "A"
    },
    {
        "type": "PG",
        "cat": "SYSTEM MONITORING",
        "title": "Pengecekan Listening Service (ss -tulpn)",
        "q": "Perintah modern pada Debian 12 untuk memeriksa seluruh port TCP/UDP yang sedang aktif mendengarkan koneksi (listening) beserta nama aplikasinya adalah...",
        "lines": [
            "root@debian12:~# ss -tulpn",
            "Netid  State   Local Address:Port  Process",
            "tcp    LISTEN  0.0.0.0:22          (\"sshd\")",
            "tcp    LISTEN  0.0.0.0:80          (\"apache2\")",
            "tcp    LISTEN  192.168.10.1:53     (\"named\")"
        ],
        "options": [
            "A. ss -tulpn atau netstat -tulpn",
            "B. ping localhost -c 4",
            "C. traceroute 127.0.0.1",
            "D. ip route show",
            "E. cat /proc/cpuinfo"
        ],
        "kunci": "A"
    },
    {
        "type": "PG",
        "cat": "DNS SERVER (BIND9)",
        "title": "Forward Zone File (/etc/bind/db.pasundan)",
        "q": "Perhatikan cuplikan zone file BIND9 berikut:\nRecord DNS yang berfungsi memetakan nama alias www.pasundan.sch.id merujuk ke domain utama pasundan.sch.id adalah...",
        "lines": [
            "@       IN      NS      pasundan.sch.id.",
            "@       IN      A       192.168.10.1",
            "www     IN      CNAME   pasundan.sch.id.",
            "mail    IN      A       192.168.10.1"
        ],
        "options": [
            "A. CNAME (Canonical Name)",
            "B. A Record (Address)",
            "C. PTR Record (Pointer)",
            "D. MX Record (Mail Exchange)",
            "E. TXT Record (Text)"
        ],
        "kunci": "A"
    },
    {
        "type": "PG",
        "cat": "DNS SERVER (BIND9)",
        "title": "Reverse DNS Lookup (PTR Record)",
        "q": "Fungsi utama dari konfigurasi Reverse DNS (db.192 / in-addr.arpa) dengan PTR Record adalah...",
        "lines": [
            "; Reverse DNS db.192",
            "@       IN      NS      pasundan.sch.id.",
            "1       IN      PTR     pasundan.sch.id.",
            "1       IN      PTR     www.pasundan.sch.id."
        ],
        "options": [
            "A. Menerjemahkan IP Address (192.168.10.1) menjadi Nama Domain",
            "B. Menerjemahkan Nama Domain menjadi IP Address",
            "C. Mengalihkan port 80 ke 443",
            "D. Memblokir serangan brute force DNS",
            "E. Mengatur alokasi IP DHCP klien"
        ],
        "kunci": "A"
    },
    {
        "type": "PG",
        "cat": "WEB SERVER (APACHE2)",
        "title": "DocumentRoot VirtualHost Apache",
        "q": "Perhatikan konfigurasi Virtual Host Apache2 di /etc/apache2/sites-available/pasundan.conf berikut:\nDirektori penyimpanan file website yang akan dimuat saat domain diakses adalah...",
        "lines": [
            "<VirtualHost *:80>",
            "    ServerName pasundan.sch.id",
            "    ServerAlias www.pasundan.sch.id",
            "    DocumentRoot /var/www/pasundan",
            "</VirtualHost>"
        ],
        "options": [
            "A. /var/www/pasundan",
            "B. /etc/apache2/sites-available",
            "C. /var/log/apache2",
            "D. /usr/share/apache2",
            "E. /etc/apache2/mods-enabled"
        ],
        "kunci": "A"
    },
    {
        "type": "PG",
        "cat": "WEB SERVER (APACHE2)",
        "title": "Aktivasi Site Apache (a2ensite)",
        "q": "Perintah CLI untuk mengaktifkan konfigurasi Virtual Host pasundan.conf pada Apache2 di Debian 12 adalah...",
        "lines": [
            "root@debian12:~# a2ensite pasundan.conf",
            "Enabling site pasundan.",
            "To activate the new configuration: systemctl reload apache2"
        ],
        "options": [
            "A. a2ensite pasundan.conf",
            "B. a2dissite pasundan.conf",
            "C. a2enmod rewrite",
            "D. apache2ctl start site",
            "E. apt install pasundan.conf"
        ],
        "kunci": "A"
    },
    {
        "type": "PG",
        "cat": "DHCP SERVER",
        "title": "Menentukan Interface DHCP (/etc/default/isc-dhcp-server)",
        "q": "Pada Debian 12, file konfigurasi yang digunakan untuk menentukan interface kartu jaringan mana yang akan menyiarkan layanan DHCP (misal: ens18) adalah...",
        "lines": [
            "root@debian12:~# cat /etc/default/isc-dhcp-server",
            "INTERFACESv4=\"ens18\"",
            "INTERFACESv6=\"\""
        ],
        "options": [
            "A. /etc/default/isc-dhcp-server",
            "B. /etc/dhcp/dhcpd.conf",
            "C. /etc/network/interfaces",
            "D. /etc/resolv.conf",
            "E. /var/lib/dhcp/dhcpd.leases"
        ],
        "kunci": "A"
    },
    {
        "type": "PG",
        "cat": "DHCP SERVER",
        "title": "Rentang Alokasi IP (Range DHCP)",
        "q": "Perhatikan konfigurasi file /etc/dhcp/dhcpd.conf berikut:\nRentang IP address yang akan dibagikan secara otomatis kepada klien adalah...",
        "lines": [
            "subnet 192.168.10.0 netmask 255.255.255.0 {",
            "    range 192.168.10.50 192.168.10.100;",
            "    option routers 192.168.10.1;",
            "}"
        ],
        "options": [
            "A. 192.168.10.50 sampai 192.168.10.100 (51 Host)",
            "B. 192.168.10.1 sampai 192.168.10.254",
            "C. Hanya IP 192.168.10.1",
            "D. 192.168.10.0 sampai 192.168.10.255",
            "E. 8.8.8.8 sampai 8.8.4.4"
        ],
        "kunci": "A"
    },
    {
        "type": "PG",
        "cat": "FTP SERVER (VSFTPD)",
        "title": "Penguncian Direktori Home (Chroot)",
        "q": "Pada konfigurasi /etc/vsftpd.conf, opsi yang harus diaktifkan agar user FTP yang login terkunci di dalam direktori home-nya masing-masing dan tidak dapat menjelajah ke direktori root sistem adalah...",
        "lines": [
            "root@debian12:~# nano /etc/vsftpd.conf",
            "anonymous_enable=NO",
            "local_enable=YES",
            "write_enable=YES",
            "chroot_local_user=YES"
        ],
        "options": [
            "A. chroot_local_user=YES",
            "B. anonymous_enable=YES",
            "C. write_enable=NO",
            "D. local_enable=NO",
            "E. listen_port=2121"
        ],
        "kunci": "A"
    },
    {
        "type": "PG",
        "cat": "FTP SERVER",
        "title": "Port Standar FTP (Control & Data)",
        "q": "Protokol FTP beroperasi menggunakan dua port TCP utama, yaitu port untuk jalur perintah (command/control) dan port untuk transfer data. Kedua port tersebut adalah...",
        "lines": [
            "FTP Control (Command channel) : Port 21 / TCP",
            "FTP Data (Data transfer)       : Port 20 / TCP"
        ],
        "options": [
            "A. Port 21 (Control) dan Port 20 (Data)",
            "B. Port 22 (Control) dan Port 23 (Data)",
            "C. Port 80 (Control) dan Port 443 (Data)",
            "D. Port 25 (Control) dan Port 110 (Data)",
            "E. Port 53 (Control) dan Port 67 (Data)"
        ],
        "kunci": "A"
    },
    {
        "type": "PG",
        "cat": "MAIL SERVER (POSTFIX)",
        "title": "Format Penyimpanan Email (Maildir/)",
        "q": "Konfigurasi 'home_mailbox = Maildir/' pada Postfix memiliki keunggulan dibandingkan format mbox, yaitu...",
        "lines": [
            "Format Mbox    : Semua email ditumpuk jadi 1 file (/var/mail/user)",
            "Format Maildir : Setiap email disimpan sebagai file terpisah di folder ~/Maildir/"
        ],
        "options": [
            "A. Setiap pesan email disimpan sebagai file tersendiri di folder user sehingga lebih aman dari korupsi data",
            "B. Menghapus seluruh email secara otomatis setelah 7 hari",
            "C. Mengompresi file email menjadi file ZIP",
            "D. Mengubah protokol SMTP menjadi HTTP",
            "E. Tidak membutuhkan kapasitas harddisk"
        ],
        "kunci": "A"
    },
    {
        "type": "PG",
        "cat": "MAIL SERVER (DOVECOT)",
        "title": "Karakteristik Protokol IMAP",
        "q": "Dovecot menyediakan protokol IMAP dan POP3 bagi klien untuk mengakses email. Karakteristik utama dari protokol IMAP adalah...",
        "lines": [
            "IMAP (Port 143 / SSL 993) : Email tersinkronisasi dan tetap tersimpan di server",
            "POP3 (Port 110 / SSL 995) : Email diunduh ke klien dan dihapus dari server"
        ],
        "options": [
            "A. Email tersinkronisasi dua arah dan tetap tersimpan di server",
            "B. Email langsung dihapus permanen dari server setelah diunduh",
            "C. Hanya bisa mengirim email tanpa bisa menerima",
            "D. Tidak mendukung koneksi terenkripsi SSL",
            "E. Menggunakan port 25"
        ],
        "kunci": "A"
    },

    # =========================================================================
    # BAGIAN 2: PILIHAN GANDA KOMPLEKS (Checklist Multi-Jawaban Benar) (Soal 16 - 25)
    # =========================================================================
    {
        "type": "PG_KOMPLEKS",
        "cat": "VIRTUALBOX & NETWORKING",
        "title": "Karakteristik Mode Adapter VirtualBox",
        "q": "Perhatikan skema topologi adapter jaringan pada VirtualBox berikut. Manakah pernyataan yang BENAR mengenai perbedaan mode jaringan VirtualBox? (Pilih semua jawaban yang benar)",
        "lines": [
            "Adapter Mode 1: NAT -> Guest dapat akses internet via IP Host",
            "Adapter Mode 2: Bridged -> Guest mendapatkan IP satu segmen dengan LAN fisik",
            "Adapter Mode 3: Host-Only -> Komunikasi tertutup antara Guest dan Host saja"
        ],
        "options": [
            "A. Mode Bridged membuat VM seolah-olah terhubung langsung ke switch/router fisik yang sama dengan Host",
            "B. Mode Host-Only memungkinkan akses langsung ke seluruh internet publik tanpa router",
            "C. Mode NAT menggunakan IP forwarding dari Host untuk koneksi keluar",
            "D. Mode Internal Network mengisolasi komunikasi hanya antar VM di host yang sama tanpa melibatkan OS Host"
        ],
        "kunci": "A, C, D"
    },
    {
        "type": "PG_KOMPLEKS",
        "cat": "DNS SERVER (BIND9)",
        "title": "Jenis-jenis DNS Record Penting",
        "q": "Dalam konfigurasi master zone file BIND9 (/etc/bind/db.domain), record manakah yang digunakan untuk memetakan domain ke IP Address dan layanan email? (Pilih semua yang benar)",
        "lines": [
            "; Master Zone Definitions",
            "@    IN   A     192.168.10.1",
            "mail IN   A     192.168.10.1",
            "@    IN   MX 10 mail.pasundan.sch.id."
        ],
        "options": [
            "A. A Record untuk memetakan nama hostname ke alamat IPv4",
            "B. MX Record untuk menentukan server penerima email (Mail Exchanger)",
            "C. CNAME Record untuk membuat nama alias dari nama domain kanonikal",
            "D. DHCP Record untuk menyewakan IP address kepada klien Windows"
        ],
        "kunci": "A, B, C"
    },
    {
        "type": "PG_KOMPLEKS",
        "cat": "DNS SERVER",
        "title": "Perintah Diagnostik & Verifikasi DNS",
        "q": "Perintah CLI apa sajakah yang dapat digunakan pada Debian 12 untuk menguji resolusi nama domain dan memvalidasi sintaks konfigurasi BIND9? (Pilih semua yang benar)",
        "lines": [
            "root@debian12:~# nslookup pasundan.sch.id",
            "root@debian12:~# dig pasundan.sch.id @127.0.0.1",
            "root@debian12:~# named-checkconf"
        ],
        "options": [
            "A. nslookup",
            "B. dig",
            "C. named-checkconf",
            "D. a2ensite"
        ],
        "kunci": "A, B, C"
    },
    {
        "type": "PG_KOMPLEKS",
        "cat": "WEB SERVER (APACHE2)",
        "title": "Modul Populer & Perintah Manajemen Apache2",
        "q": "Manakah perintah dan modul Apache2 berikut yang BENAR digunakan pada Debian 12 untuk mengelola modul dan website? (Pilih semua yang benar)",
        "lines": [
            "root@debian12:~# a2enmod ssl",
            "root@debian12:~# a2enmod rewrite",
            "root@debian12:~# a2ensite 000-default.conf",
            "root@debian12:~# a2dissite old-web.conf"
        ],
        "options": [
            "A. a2enmod ssl untuk mengaktifkan dukungan enkripsi HTTPS",
            "B. a2enmod rewrite untuk mengaktifkan manipulasi/pengalihan URL (.htaccess)",
            "C. a2ensite untuk mengaktifkan file konfigurasi virtual host dari sites-available",
            "D. a2enmod apache3 untuk mengupgrade versi web server secara otomatis"
        ],
        "kunci": "A, B, C"
    },
    {
        "type": "PG_KOMPLEKS",
        "cat": "WEB SERVER",
        "title": "Keamanan Web Server & HTTPS",
        "q": "Langkah-langkah apa sajakah yang wajib dilakukan saat mengkonfigurasi HTTPS (SSL/TLS) pada web server Apache2 di Debian 12? (Pilih semua yang benar)",
        "lines": [
            "1. Generate / Pasang Sertifikat SSL (.crt) & Private Key (.key)",
            "2. Aktifkan modul SSL (a2enmod ssl)",
            "3. Konfigurasi VirtualHost pada Port 443 (SSLEngine on, SSLCertificateFile, SSLCertificateKeyFile)"
        ],
        "options": [
            "A. Membuka dan mendengarkan port 443 di /etc/apache2/ports.conf",
            "B. Menentukan path SSLCertificateFile dan SSLCertificateKeyFile pada VirtualHost",
            "C. Mengaktifkan direktif SSLEngine on",
            "D. Menghapus user www-data dari sistem"
        ],
        "kunci": "A, B, C"
    },
    {
        "type": "PG_KOMPLEKS",
        "cat": "DHCP SERVER",
        "title": "Parameter Wajib dalam Subnet DHCP",
        "q": "Di dalam deklarasi blok 'subnet' pada file /etc/dhcp/dhcpd.conf, parameter konfigurasi manakah yang lazim disertakan untuk memberikan konfigurasi jaringan lengkap kepada klien? (Pilih semua yang benar)",
        "lines": [
            "subnet 192.168.10.0 netmask 255.255.255.0 {",
            "    range 192.168.10.100 192.168.10.200;",
            "    option routers 192.168.10.1;",
            "    option domain-name-servers 192.168.10.1, 8.8.8.8;",
            "}"
        ],
        "options": [
            "A. range (rentang IP awal hingga IP akhir yang disewakan)",
            "B. option routers (alamat default gateway klien)",
            "C. option domain-name-servers (alamat DNS resolver klien)",
            "D. option anonymous_enable (untuk membuka akses tanpa password)"
        ],
        "kunci": "A, B, C"
    },
    {
        "type": "PG_KOMPLEKS",
        "cat": "FTP SERVER (VSFTPD)",
        "title": "Opsi Keamanan & Izin pada VSFTPD",
        "q": "Opsi konfigurasi manakah pada file /etc/vsftpd.conf yang digunakan untuk mengatur perizinan login user lokal dan hak unggah file? (Pilih semua yang benar)",
        "lines": [
            "anonymous_enable=NO",
            "local_enable=YES",
            "write_enable=YES",
            "chroot_local_user=YES"
        ],
        "options": [
            "A. local_enable=YES (mengizinkan user sistem Linux login ke FTP)",
            "B. write_enable=YES (mengizinkan perintah penulisan / upload file)",
            "C. anonymous_enable=NO (menonaktifkan akses anonim tanpa autentikasi)",
            "D. port_forwarding=ENABLED (mengubah port 21 menjadi port 80)"
        ],
        "kunci": "A, B, C"
    },
    {
        "type": "PG_KOMPLEKS",
        "cat": "MAIL SERVER",
        "title": "Komponen Arsitektur Mail Server",
        "q": "Sebuah Mail Server lengkap pada Linux umumnya terdiri atas beberapa komponen perangkat lunak dengan peran spesifik. Manakah pasangan komponen dan peran yang BENAR? (Pilih semua yang benar)",
        "lines": [
            "MTA (Mail Transfer Agent) : Postfix (SMTP - Kirim Email)",
            "MDA (Mail Delivery Agent) : Dovecot (IMAP/POP3 - Ambil Email)",
            "MUA (Mail User Agent)     : Roundcube / Thunderbird (Antarmuka Pengguna)"
        ],
        "options": [
            "A. Postfix bertindak sebagai MTA (Mail Transfer Agent) yang menangani protokol SMTP",
            "B. Dovecot bertindak sebagai MDA/Server Akses yang menyediakan protokol IMAP dan POP3",
            "C. Roundcube bertindak sebagai Webmail Client (MUA) yang diakses pengguna melalui web browser",
            "D. BIND9 bertindak sebagai Mail Transfer Agent utama pengganti Postfix"
        ],
        "kunci": "A, B, C"
    },
    {
        "type": "PG_KOMPLEKS",
        "cat": "MAIL SERVER",
        "title": "Port Standar Layanan Email",
        "q": "Manakah nomor port dan protokol email berikut yang BENAR dan berstandar internasional? (Pilih semua yang benar)",
        "lines": [
            "SMTP Plain: Port 25 / TCP",
            "POP3 Plain: Port 110 / TCP",
            "IMAP Plain: Port 143 / TCP",
            "IMAPS (SSL): Port 993 / TCP"
        ],
        "options": [
            "A. Port 25 (SMTP untuk pengiriman email antar server)",
            "B. Port 110 (POP3 untuk pengambilan email klien)",
            "C. Port 143 (IMAP untuk akses dan sinkronisasi email)",
            "D. Port 80 (SMTP terenkripsi TLS)"
        ],
        "kunci": "A, B, C"
    },
    {
        "type": "PG_KOMPLEKS",
        "cat": "LINUX SYSTEMD",
        "title": "Manajemen Service dengan systemctl",
        "q": "Perintah systemctl apa sajakah yang digunakan oleh administrator untuk mengontrol lifecycle daemon/service jaringan pada Debian 12? (Pilih semua yang benar)",
        "lines": [
            "root@debian12:~# systemctl start apache2",
            "root@debian12:~# systemctl enable named",
            "root@debian12:~# systemctl status isc-dhcp-server",
            "root@debian12:~# systemctl reload postfix"
        ],
        "options": [
            "A. systemctl start <service> untuk menjalankan service",
            "B. systemctl enable <service> agar service otomatis berjalan saat sistem booting",
            "C. systemctl status <service> untuk melihat kondisi dan log error service saat ini",
            "D. systemctl format <service> untuk mereset seluruh harddisk server"
        ],
        "kunci": "A, B, C"
    },

    # =========================================================================
    # BAGIAN 3: BENAR / SALAH (TRUE / FALSE) (Soal 26 - 30)
    # =========================================================================
    {
        "type": "BENAR_SALAH",
        "cat": "VIRTUALBOX",
        "title": "Mode Promiscuous pada Adapter VirtualBox",
        "q": "Pernyataan: 'Mengaktifkan Promiscuous Mode: Allow All pada VirtualBox diperlukan agar VM dapat menangkap dan memproses semua paket jaringan termasuk paket untuk sub-interface VLAN atau MAC address virtual di dalam VM.'",
        "lines": [
            "VirtualBox Network Settings:",
            "Promiscuous Mode: [ Allow All ]",
            "Status Evaluasi: Pernyataan Konseptual Jaringan Virtual"
        ],
        "options": [
            "A. Benar",
            "B. Salah"
        ],
        "kunci": "BENAR"
    },
    {
        "type": "BENAR_SALAH",
        "cat": "DNS SERVER (BIND9)",
        "title": "Penyelesaian Alamat IP ke Nama Domain",
        "q": "Pernyataan: 'Fungsi dari Forward Lookup Zone pada BIND9 adalah menerjemahkan IP Address menjadi Nama Domain, sedangkan Reverse Lookup Zone menerjemahkan Nama Domain menjadi IP Address.'",
        "lines": [
            "Klaim: Forward Zone = IP to Domain | Reverse Zone = Domain to IP",
            "Koreksi Teori: Forward = Domain to IP | Reverse = IP to Domain",
            "Status Evaluasi: Cek Kebenaran Pernyataan"
        ],
        "options": [
            "A. Benar",
            "B. Salah"
        ],
        "kunci": "SALAH"
    },
    {
        "type": "BENAR_SALAH",
        "cat": "WEB SERVER (APACHE2)",
        "title": "Penggunaan File Konfigurasi .htaccess",
        "q": "Pernyataan: 'Direktif AllowOverride All di dalam blok <Directory /var/www/> pada apache2.conf diperlukan agar Apache2 dapat membaca dan menerapkan aturan per-direktori yang ditulis di dalam file .htaccess.'",
        "lines": [
            "<Directory /var/www/>",
            "    Options Indexes FollowSymLinks",
            "    AllowOverride All",
            "    Require all granted",
            "</Directory>"
        ],
        "options": [
            "A. Benar",
            "B. Salah"
        ],
        "kunci": "BENAR"
    },
    {
        "type": "BENAR_SALAH",
        "cat": "DHCP SERVER",
        "title": "DHCP Relay Agent pada Router",
        "q": "Pernyataan: 'DHCP Broadcast (DHCP DISCOVER) dari klien tidak dapat melewati router jaringan secara alami tanpa adanya konfigurasi DHCP Relay Agent (IP Helper Address) pada router tersebut.'",
        "lines": [
            "Client LAN (VLAN 10) ---> [ Router / DHCP Relay ] ---> DHCP Server (VLAN 20)",
            "Karakteristik Broadcast Layer 3: Dibatasi oleh Router Broadcast Domain",
            "Status Evaluasi: Cek Kebenaran Pernyataan"
        ],
        "options": [
            "A. Benar",
            "B. Salah"
        ],
        "kunci": "BENAR"
    },
    {
        "type": "BENAR_SALAH",
        "cat": "FTP SERVER",
        "title": "Keamanan Port FTP vs SFTP",
        "q": "Pernyataan: 'Protokol FTP standar (Port 21) mengenkripsi kredensial login (username dan password) saat dikirim melalui jaringan, sehingga tidak dapat disadap menggunakan software packet sniffer seperti Wireshark.'",
        "lines": [
            "FTP Traffic: USER siswa / PASS rahasia123 (Plaintext)",
            "Wireshark Capture: Kredensial terbaca jelas tanpa enkripsi",
            "Status Evaluasi: Cek Kebenaran Pernyataan"
        ],
        "options": [
            "A. Benar",
            "B. Salah"
        ],
        "kunci": "SALAH"
    },

    # =========================================================================
    # BAGIAN 4: MENJODOHKAN (MATCHING) (Soal 31 - 35)
    # =========================================================================
    {
        "type": "MENJODOHKAN",
        "cat": "JARINGAN & PORT",
        "title": "Menjodohkan Layanan Server dengan Nomor Port Standar",
        "q": "Jodohkanlah nama layanan server jaringan di Debian 12 sebelah kiri dengan nomor port default yang sesuai di sebelah kanan:",
        "lines": [
            "Premis (Layanan)        <=> Respon (Nomor Port)",
            "DNS Server (BIND9)      <=> Port 53",
            "Web Server HTTP         <=> Port 80",
            "FTP Control             <=> Port 21",
            "DHCP Server             <=> Port 67"
        ],
        "pairs": [
            "DNS Server (BIND9) <=> Port 53",
            "Web Server HTTP <=> Port 80",
            "FTP Control <=> Port 21",
            "DHCP Server <=> Port 67",
            "Mail Server SMTP <=> Port 25"
        ]
    },
    {
        "type": "MENJODOHKAN",
        "cat": "DEBIAN 12 CONFIG",
        "title": "Menjodohkan Layanan dengan File Konfigurasi Utamanya",
        "q": "Jodohkanlah nama layanan server di sebelah kiri dengan path file konfigurasi utamanya pada sistem operasi Debian 12 di sebelah kanan:",
        "lines": [
            "Layanan Server          <=> File Konfigurasi Utama",
            "DNS Master Zone         <=> /etc/bind/named.conf.local",
            "Virtual Host Apache     <=> /etc/apache2/sites-available/000-default.conf",
            "Layanan DHCP            <=> /etc/dhcp/dhcpd.conf",
            "Layanan FTP (VSFTPD)    <=> /etc/vsftpd.conf"
        ],
        "pairs": [
            "DNS Master Zone <=> /etc/bind/named.conf.local",
            "Virtual Host Apache <=> /etc/apache2/sites-available/000-default.conf",
            "Layanan DHCP <=> /etc/dhcp/dhcpd.conf",
            "Layanan FTP (VSFTPD) <=> /etc/vsftpd.conf",
            "Mail Server (Postfix) <=> /etc/postfix/main.cf"
        ]
    },
    {
        "type": "MENJODOHKAN",
        "cat": "DNS RECORDS",
        "title": "Menjodohkan Tipe DNS Record dengan Fungsinya",
        "q": "Jodohkanlah tipe DNS Record di sebelah kiri dengan fungsinya yang tepat di sebelah kanan:",
        "lines": [
            "Tipe DNS Record         <=> Fungsi Record",
            "A Record                <=> Memetakan Hostname ke IPv4",
            "AAAA Record             <=> Memetakan Hostname ke IPv6",
            "CNAME Record            <=> Membuat Nama Alias Hostname",
            "PTR Record              <=> Memetakan IP ke Hostname (Reverse)"
        ],
        "pairs": [
            "A Record <=> Memetakan Hostname ke IPv4",
            "AAAA Record <=> Memetakan Hostname ke IPv6",
            "CNAME Record <=> Membuat Nama Alias Hostname",
            "PTR Record <=> Memetakan IP ke Hostname (Reverse DNS)",
            "MX Record <=> Menentukan Mail Server Penerima Email"
        ]
    },
    {
        "type": "MENJODOHKAN",
        "cat": "CLI COMMANDS",
        "title": "Menjodohkan Perintah CLI Debian dengan Tujuannya",
        "q": "Jodohkanlah perintah baris perintah (CLI) di sebelah kiri dengan fungsi tujuannya di sebelah kanan:",
        "lines": [
            "Perintah CLI            <=> Fungsi & Tujuan",
            "a2ensite web.conf       <=> Mengaktifkan Virtual Host Apache",
            "named-checkconf         <=> Memeriksa Sintaks Konfigurasi BIND9",
            "tail -f error.log       <=> Memantau Log Error Real-Time",
            "chmod 755 /var/www      <=> Mengatur Izin Hak Akses Direktori"
        ],
        "pairs": [
            "a2ensite web.conf <=> Mengaktifkan Virtual Host Apache",
            "named-checkconf <=> Memeriksa Sintaks Konfigurasi BIND9",
            "tail -f error.log <=> Memantau Log Error Real-Time",
            "chmod 755 /var/www <=> Mengatur Izin Hak Akses Direktori",
            "systemctl restart named <=> Memuat Ulang Layanan DNS BIND9"
        ]
    },
    {
        "type": "MENJODOHKAN",
        "cat": "MAIL PROTOCOLS",
        "title": "Menjodohkan Protokol Email dengan Karakteristiknya",
        "q": "Jodohkanlah protokol email di sebelah kiri dengan karakteristik tugasnya di sebelah kanan:",
        "lines": [
            "Protokol Email          <=> Karakteristik & Peran",
            "SMTP (Port 25)          <=> Mengirim dan Menyalurkan Email Antar Server",
            "POP3 (Port 110)         <=> Mengunduh Email ke Klien dan Menghapus dari Server",
            "IMAP (Port 143)         <=> Sinkronisasi Dua Arah Email di Server",
            "IMAPS (Port 993)        <=> Akses IMAP Aman Terenkripsi SSL/TLS"
        ],
        "pairs": [
            "SMTP (Port 25) <=> Mengirim dan Menyalurkan Email Antar Server",
            "POP3 (Port 110) <=> Mengunduh Email ke Klien dan Menghapus dari Server",
            "IMAP (Port 143) <=> Sinkronisasi Dua Arah Email Tetap di Server",
            "IMAPS (Port 993) <=> Akses IMAP Aman Terenkripsi SSL/TLS",
            "Roundcube <=> Webmail Client Antarmuka Browser"
        ]
    },

    # =========================================================================
    # BAGIAN 5: ESAI / URAIAN STUDI KASUS (Soal 36 - 40)
    # =========================================================================
    {
        "type": "ESAI",
        "cat": "TROUBLESHOOTING DNS",
        "title": "Studi Kasus Kegagalan Service BIND9",
        "q": "Jelaskan langkah-langkah troubleshooting sistematis yang harus Anda lakukan jika setelah mengedit file db.pasundan dan menjalankan 'systemctl restart named', status service BIND9 mengalami status FAILED (Error)!",
        "lines": [
            "root@debian12:~# systemctl status named",
            "● named.service - BIND Domain Name Server",
            "   Active: failed (Result: exit-code) since Sat 2026-08-29 07:10:00",
            "   Process: 4120 ExecStart=/usr/sbin/named (code=exited, status=1/FAILURE)"
        ],
        "kunci": "1. Periksa log detail menggunakan perintah 'journalctl -u named -e' atau 'named-checkconf' dan 'named-checkzone pasundan.sch.id /etc/bind/db.pasundan'.\n2. Telusuri nomor baris yang dilaporkan mengalami kesalahan sintaks (misal: titik '.' yang tertinggal di akhir FQDN, nomor serial SOA yang tidak valid, atau tanda kurung kurawal yang belum tertutup).\n3. Perbaiki baris yang error menggunakan text editor (nano/vim).\n4. Jalankan kembali 'named-checkconf' hingga tidak ada pesan error.\n5. Restart service dengan 'systemctl restart named' dan verifikasi status dengan 'systemctl status named'."
    },
    {
        "type": "ESAI",
        "cat": "VIRTUALHOST APACHE",
        "title": "Konfigurasi Multi VirtualHost pada Satu IP",
        "q": "Jelaskan konsep dan langkah-langkah konfigurasi Name-Based Virtual Hosting pada Apache2 di Debian 12 sehingga satu server dengan IP 192.168.10.1 dapat melayani dua website berbeda, yaitu 'smk.sch.id' (/var/www/smk) dan 'cbt.sch.id' (/var/www/cbt)!",
        "lines": [
            "IP Server Tunggal : 192.168.10.1",
            "Website 1 : smk.sch.id -> DocumentRoot /var/www/smk",
            "Website 2 : cbt.sch.id -> DocumentRoot /var/www/cbt"
        ],
        "kunci": "1. Buat direktori root untuk kedua web: 'mkdir -p /var/www/smk /var/www/cbt' dan isi file index.html masing-masing.\n2. Berikan hak akses: 'chown -R www-data:www-data /var/www/' dan 'chmod -R 755 /var/www/'.\n3. Buat file virtualhost pertama '/etc/apache2/sites-available/smk.conf' dengan ServerName smk.sch.id dan DocumentRoot /var/www/smk.\n4. Buat file virtualhost kedua '/etc/apache2/sites-available/cbt.conf' dengan ServerName cbt.sch.id dan DocumentRoot /var/www/cbt.\n5. Aktifkan kedua situs: 'a2ensite smk.conf' dan 'a2ensite cbt.conf'.\n6. Reload konfigurasi web server: 'systemctl reload apache2'."
    },
    {
        "type": "ESAI",
        "cat": "DHCP RESERVATION",
        "title": "Konfigurasi Static Lease (MAC Binding) pada DHCP",
        "q": "Tuliskan cuplikan blok konfigurasi pada file /etc/dhcp/dhcpd.conf untuk melakukan reservasi IP (Static Lease) agar komputer Server Ujian dengan MAC Address '08:00:27:AA:BB:CC' selalu memperoleh IP Address statis '192.168.10.200' beserta penjelasannya!",
        "lines": [
            "Nama Host   : Server-Ujian",
            "MAC Address : 08:00:27:AA:BB:CC",
            "Reserved IP : 192.168.10.200"
        ],
        "kunci": "host Server-Ujian {\n    hardware ethernet 08:00:27:AA:BB:CC;\n    fixed-address 192.168.10.200;\n}\nPenjelasan: Direktif 'hardware ethernet' memetakan alamat fisik kartu jaringan (MAC Address) klien, dan direktif 'fixed-address' mengunci pemberian IP tertentu agar DHCP server tidak memberikan IP tersebut kepada perangkat lain."
    },
    {
        "type": "ESAI",
        "cat": "FTP SECURITY",
        "title": "Pengamanan Akses FTP Server (vsftpd)",
        "q": "Sebutkan dan jelaskan minimal 3 parameter konfigurasi pada /etc/vsftpd.conf yang wajib disetting oleh administrator untuk mencegah akses ilegal dan membatasi ruang jelajah user pada FTP Server!",
        "lines": [
            "Poin Pengamanan FTP:",
            "1. Pembatasan User Anonim",
            "2. Penguncian Direktori Home User",
            "3. Pembatasan Izin Penulisan File"
        ],
        "kunci": "1. anonymous_enable=NO : Mematikan akun anonim agar tidak ada pengguna asing yang dapat masuk tanpa autentikasi username & password resmi.\n2. chroot_local_user=YES : Mengunci user yang berhasil login agar hanya berada di dalam direktori home mereka sendiri dan tidak bisa mengakses file sistem root (/etc, /bin, dll).\n3. allow_writeable_chroot=YES : Mengizinkan user mengunggah file ke direktori chroot dengan aman tanpa menyebabkan konflik keamanan VSFTPD.\n4. write_enable=YES : Mengontrol apakah user diizinkan mengunggah/memodifikasi file atau hanya bersifat read-only."
    },
    {
        "type": "ESAI",
        "cat": "ALUR INTEGRASI MAIL",
        "title": "Alur Pengiriman Email dari Webmail hingga Inbox Klien",
        "q": "Uraikan alur kerja interaksi perangkat lunak (Roundcube, Postfix, Dovecot, dan DNS MX) saat seorang pengguna 'guru@pasundan.sch.id' mengirimkan pesan email kepada 'siswa@pasundan.sch.id' melalui webmail browser!",
        "lines": [
            "User Browser (Roundcube) ---> Postfix (SMTP Port 25)",
            "Postfix ---> Local Maildir Delivery (~/Maildir/new/)",
            "Penerima (Dovecot IMAP Port 143) ---> Roundcube Inbox Siswa"
        ],
        "kunci": "1. Guru membuat pesan di browser via Webmail Roundcube dan menekan tombol 'Kirim'.\n2. Roundcube meneruskan pesan tersebut ke Postfix (MTA) melalui protokol SMTP (Port 25/587).\n3. Postfix membaca alamat tujuan 'siswa@pasundan.sch.id', memeriksa konfigurasi 'mydestination' (domain lokal), dan mengecek DNS MX record domain.\n4. Karena domain lokal, Postfix menyerahkan email ke direktori penyimpanan kotak surat siswa di '/home/siswa/Maildir/new/'.\n5. Siswa membuka webmail browsernya; Roundcube siswa meminta daftar email ke Dovecot via protokol IMAP (Port 143).\n6. Dovecot membaca file email baru dari direktori Maildir siswa dan menampilkannya di halaman Inbox webmail siswa."
    }
]

def generate_varied_linux_docx():
    print("🚀 Generating 40 Varied Linux & Network questions (PG, PG Kompleks, Benar/Salah, Menjodohkan, Esai)...")
    doc = Document()

    # Document Header
    title = doc.add_heading("BANK SOAL TEKNIK KOMPUTER & JARINGAN (40 BUTIR KOMBINASI)", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Topik: LINUX DEBIAN 12, VIRTUALBOX, DNS (BIND9), WEB (APACHE2), DHCP, FTP, & MAIL SERVER")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub2 = doc.add_paragraph("Bentuk Soal: Pilihan Ganda (PG), PG Kompleks (Checklist), Benar/Salah (T/F), Menjodohkan (Matching), & Esai Uraian")
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub3 = doc.add_paragraph("Dilengkapi 40 Diagram Terminal CLI & Topologi Server Beresolusi Tinggi")
    sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("=" * 80)

    for i, item in enumerate(QUESTIONS):
        q_num = i + 1
        q_type = item["type"]
        img_name = f"var_linux_q{q_num}.png"
        img_path = create_terminal_image(img_name, q_type, item["cat"], item["title"], item["lines"])

        # Paragraph for question number + tag + content
        p = doc.add_paragraph()
        p.add_run(f"{q_num}. [TIPE: {q_type}] [{item['cat']}] {item['q']}").bold = True

        # Add image
        doc.add_picture(img_path, width=Inches(5.0))

        # Handle different question types
        if q_type in ["PG", "PG_KOMPLEKS", "BENAR_SALAH"]:
            for opt in item["options"]:
                doc.add_paragraph(opt)
            p_key = doc.add_paragraph()
            p_key.add_run(f"KUNCI: {item['kunci']}").bold = True
        elif q_type == "MENJODOHKAN":
            for pair in item["pairs"]:
                doc.add_paragraph(f"PASANGAN: {pair}")
            p_key = doc.add_paragraph()
            p_key.add_run(f"KUNCI: {len(item['pairs'])} Pasangan Terhubung Benar").bold = True
        elif q_type == "ESAI":
            p_key = doc.add_paragraph()
            p_key.add_run("KUNCI / PEDOMAN PENILAIAN:").bold = True
            doc.add_paragraph(item["kunci"])

        doc.add_paragraph() # Empty separator

    # Destination paths
    out_scratch = os.path.join(os.path.dirname(__file__), "../BANK_SOAL_40_LINUX_DEBIAN12_VBOX.docx")
    out_public = os.path.join(os.path.dirname(__file__), "../public/BANK_SOAL_40_LINUX_DEBIAN12_VBOX.docx")
    out_desktop = os.path.expanduser("~/Desktop/BANK_SOAL_40_LINUX_DEBIAN12_VBOX.docx")
    out_downloads = os.path.expanduser("~/Downloads/BANK_SOAL_40_LINUX_DEBIAN12_VBOX.docx")

    doc.save(out_scratch)
    shutil.copy(out_scratch, out_public)
    try:
        shutil.copy(out_scratch, out_desktop)
        shutil.copy(out_scratch, out_downloads)
    except Exception as e:
        print(f"Copy warning: {e}")

    print(f"✅ Dokumen 40 Soal Berbagai Bentuk berhasil dibuat di:\n - {out_desktop}\n - {out_downloads}\n - {out_scratch}")

if __name__ == "__main__":
    generate_varied_linux_docx()
