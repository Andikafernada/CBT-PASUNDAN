import os
import sys
import shutil

if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageDraw, ImageFont

TEMP_DIR = os.path.join(os.path.dirname(__file__), "temp_linux_images")
os.makedirs(TEMP_DIR, exist_ok=True)

def create_terminal_image(filename, category, title, terminal_lines, bg_color="#0f172a", border_color="#38bdf8"):
    width, height = 750, 240
    img = Image.new("RGB", (width, height), color=bg_color)
    draw = ImageDraw.Draw(img)

    # Outer border
    draw.rectangle([8, 8, width - 8, height - 8], outline=border_color, width=2)

    # Terminal Top Bar
    draw.rectangle([8, 8, width - 8, 38], fill="#1e293b")
    # Red, Yellow, Green window buttons
    draw.ellipse([20, 18, 30, 28], fill="#ef4444")
    draw.ellipse([36, 18, 46, 28], fill="#f59e0b")
    draw.ellipse([52, 18, 62, 28], fill="#10b981")
    draw.text((80, 15), f"root@debian12:~# [{category}] - {title}", fill="#94a3b8")

    # Terminal Code Lines
    y = 52
    for line in terminal_lines:
        if line.startswith("root@") or line.startswith("$") or line.startswith("#"):
            draw.text((25, y), line, fill="#38bdf8")
        elif "OK" in line or "active (running)" in line or "SUCCESS" in line:
            draw.text((25, y), line, fill="#4ade80")
        elif "ERROR" in line or "FAILED" in line or "FAIL" in line:
            draw.text((25, y), line, fill="#f87171")
        elif line.startswith(";"):
            draw.text((25, y), line, fill="#64748b")
        else:
            draw.text((25, y), line, fill="#f8fafc")
        y += 26

    path = os.path.join(TEMP_DIR, filename)
    img.save(path, format="PNG")
    return path

# 40 Complete Questions on Linux, DNS, Web Apache2, DHCP, FTP, Mail on Debian 12 & VirtualBox
QUESTIONS = [
    # 1 - 6: Linux Dasar & VirtualBox Network Mode
    {
        "cat": "VIRTUALBOX & LINUX",
        "title": "Mode Jaringan VirtualBox (Host-Only vs Bridged)",
        "q": "Pada VirtualBox, administrator ingin agar VM Debian 12 dapat berkomunikasi dua arah dengan OS Host (Windows) tanpa terhubung ke internet luar, serta memiliki subnet IP statis tersendiri. Mode adapter jaringan VirtualBox yang paling tepat digunakan adalah...",
        "lines": [
            "Settings > Network > Adapter 1",
            "Attached to: [ Host-only Adapter ]",
            "Name: VirtualBox Host-Only Ethernet Adapter",
            "Promiscuous Mode: Allow All",
            "Cable Connected: [✓] Checked"
        ],
        "a": "A. Host-only Adapter",
        "b": "B. NAT (Network Address Translation)",
        "c": "C. Bridged Adapter",
        "d": "D. Internal Network",
        "e": "E. Generic Driver",
        "kunci": "A"
    },
    {
        "cat": "VIRTUALBOX",
        "title": "VirtualBox Port Forwarding pada NAT",
        "q": "Jika VM Debian 12 menggunakan adapter jaringan NAT pada VirtualBox dan menjalankan web server Apache di port 80, konfigurasi Port Forwarding pada VirtualBox agar web server dapat diakses dari browser Host melalui alamat http://localhost:8080 adalah...",
        "lines": [
            "Rule 1: Name: HTTP-Web, Protocol: TCP",
            "Host IP: 127.0.0.1, Host Port: 8080",
            "Guest IP: 10.0.2.15, Guest Port: 80",
            "Status: Active Forwarding"
        ],
        "a": "A. Host Port: 8080, Guest Port: 80",
        "b": "B. Host Port: 80, Guest Port: 8080",
        "c": "C. Host Port: 22, Guest Port: 80",
        "d": "D. Host Port: 443, Guest Port: 8080",
        "e": "E. Host Port: 53, Guest Port: 53",
        "kunci": "A"
    },
    {
        "cat": "DEBIAN 12 CLI",
        "title": "Konfigurasi IP Statis Debian 12 (/etc/network/interfaces)",
        "q": "Perhatikan konfigurasi file /etc/network/interfaces pada Debian 12 berikut:\nPerintah yang tepat untuk menerapkan perubahan konfigurasi IP address tersebut tanpa melakukan restart server secara keseluruhan adalah...",
        "lines": [
            "auto ens18",
            "iface ens18 inet static",
            "    address 192.168.10.1/24",
            "    gateway 192.168.10.254",
            "    dns-nameservers 192.168.10.1 8.8.8.8"
        ],
        "a": "A. systemctl restart networking",
        "b": "B. systemctl restart apache2",
        "c": "C. ifconfig ens18 restart",
        "d": "D. ip address reload ens18",
        "e": "E. service network-manager stop",
        "kunci": "A"
    },
    {
        "cat": "DEBIAN 12 CLI",
        "title": "Manajemen Repositori & Paket (APT)",
        "q": "Pada Debian 12 (Bookworm), file utama yang digunakan untuk mendaftarkan URL repositori paket instalasi software berada pada direktori...",
        "lines": [
            "root@debian12:~# cat /etc/apt/sources.list",
            "deb http://deb.debian.org/debian bookworm main contrib non-free",
            "deb http://security.debian.org/debian-security bookworm-security main",
            "root@debian12:~# apt update"
        ],
        "a": "A. /etc/apt/sources.list",
        "b": "B. /etc/network/interfaces",
        "c": "C. /etc/environment",
        "d": "D. /var/cache/apt/archives",
        "e": "E. /etc/resolv.conf",
        "kunci": "A"
    },
    {
        "cat": "LINUX PERMISSIONS",
        "title": "Hak Akses File & Direktori (chmod & chown)",
        "q": "Administrator ingin memberikan hak akses penuh (baca, tulis, eksekusi) untuk pemilik file (owner), serta hanya hak baca dan eksekusi untuk grup dan user lainnya pada folder /var/www/html. Nilai numerik chmod yang tepat adalah...",
        "lines": [
            "root@debian12:~# chmod 755 -R /var/www/html",
            "root@debian12:~# chown -R www-data:www-data /var/www/html",
            "root@debian12:~# ls -ld /var/www/html",
            "drwxr-xr-x 2 www-data www-data 4096 Aug 29 07:00 /var/www/html"
        ],
        "a": "A. chmod 755 /var/www/html",
        "b": "B. chmod 777 /var/www/html",
        "c": "C. chmod 644 /var/www/html",
        "d": "D. chmod 700 /var/www/html",
        "e": "E. chmod 600 /var/www/html",
        "kunci": "A"
    },
    {
        "cat": "SYSTEMD & MONITORING",
        "title": "Pengecekan Port Aktif & Listening Service",
        "q": "Perintah modern pada Debian 12 untuk memeriksa seluruh port TCP yang sedang aktif mendengarkan koneksi (listening) beserta nama aplikasinya adalah...",
        "lines": [
            "root@debian12:~# ss -tulpn",
            "Netid  State   Recv-Q  Send-Q  Local Address:Port  Process",
            "tcp    LISTEN  0       128     0.0.0.0:22          (\"sshd\")",
            "tcp    LISTEN  0       511     0.0.0.0:80          (\"apache2\")",
            "tcp    LISTEN  0       128     192.168.10.1:53     (\"named\")"
        ],
        "a": "A. ss -tulpn atau netstat -tulpn",
        "b": "B. ping localhost -c 4",
        "c": "C. traceroute 127.0.0.1",
        "d": "D. ip route show",
        "e": "E. cat /proc/cpuinfo",
        "kunci": "A"
    },

    # 7 - 14: DNS Server (BIND9)
    {
        "cat": "DNS SERVER (BIND9)",
        "title": "Konfigurasi Zone File (/etc/bind/named.conf.local)",
        "q": "Perhatikan cuplikan konfigurasi zone forward BIND9 pada file /etc/bind/named.conf.local berikut:\nFile database yang memetakan nama domain pasundan.sch.id menjadi IP address disimpan pada file...",
        "lines": [
            "zone \"pasundan.sch.id\" {",
            "    type master;",
            "    file \"/etc/bind/db.pasundan\";",
            "};",
            "zone \"10.168.192.in-addr.arpa\" {",
            "    type master;",
            "    file \"/etc/bind/db.192\";",
            "};"
        ],
        "a": "A. /etc/bind/db.pasundan",
        "b": "B. /etc/bind/named.conf.options",
        "c": "C. /etc/resolv.conf",
        "d": "D. /etc/hosts",
        "e": "E. /etc/bind/db.192",
        "kunci": "A"
    },
    {
        "cat": "DNS SERVER (BIND9)",
        "title": "Forwarding DNS Options (/etc/bind/named.conf.options)",
        "q": "Agar DNS server lokal BIND9 dapat meneruskan (forward) query domain internet ke DNS Google saat klien membuka google.com, konfigurasi yang harus diaktifkan pada named.conf.options adalah...",
        "lines": [
            "options {",
            "    directory \"/var/cache/bind\";",
            "    forwarders {",
            "        8.8.8.8;",
            "        8.8.4.4;",
            "    };",
            "    dnssec-validation no;",
            "};"
        ],
        "a": "A. forwarders { 8.8.8.8; 8.8.4.4; };",
        "b": "B. listen-on port 80 { any; };",
        "c": "C. allow-transfer { none; };",
        "d": "D. zone-statistics yes;",
        "e": "E. recursion no;",
        "kunci": "A"
    },
    {
        "cat": "DNS SERVER (BIND9)",
        "title": "Tipe Record DNS (A, CNAME, MX, PTR)",
        "q": "Pada file zone forward /etc/bind/db.pasundan, record yang digunakan untuk membuat nama alias seperti www.pasundan.sch.id merujuk ke domain utama pasundan.sch.id adalah...",
        "lines": [
            "@       IN      NS      pasundan.sch.id.",
            "@       IN      A       192.168.10.1",
            "www     IN      CNAME   pasundan.sch.id.",
            "mail    IN      A       192.168.10.1",
            "@       IN      MX  10  mail.pasundan.sch.id."
        ],
        "a": "A. CNAME (Canonical Name)",
        "b": "B. A Record (Address)",
        "c": "C. PTR Record (Pointer)",
        "d": "D. MX Record (Mail Exchange)",
        "e": "E. TXT Record (Text)",
        "kunci": "A"
    },
    {
        "cat": "DNS SERVER (BIND9)",
        "title": "Reverse DNS Lookup (PTR Record)",
        "q": "Fungsi utama dari konfigurasi Reverse DNS (db.192 / in-addr.arpa) dengan PTR Record adalah...",
        "lines": [
            "; Reverse DNS db.192",
            "@       IN      NS      pasundan.sch.id.",
            "1       IN      PTR     pasundan.sch.id.",
            "1       IN      PTR     www.pasundan.sch.id.",
            "1       IN      PTR     mail.pasundan.sch.id."
        ],
        "a": "A. Menerjemahkan IP Address (192.168.10.1) menjadi Nama Domain",
        "b": "B. Menerjemahkan Nama Domain menjadi IP Address",
        "c": "C. Mengalihkan port 80 ke 443",
        "d": "D. Memblokir serangan brute force DNS",
        "e": "E. Mengatur alokasi IP DHCP klien",
        "kunci": "A"
    },
    {
        "cat": "DNS DIAGNOSTICS",
        "title": "Uji Resolusi DNS (nslookup & dig)",
        "q": "Perhatikan hasil eksekusi perintah nslookup berikut:\nBerdasarkan output tersebut, status query DNS server dapat dinyatakan...",
        "lines": [
            "root@debian12:~# nslookup www.pasundan.sch.id",
            "Server:         192.168.10.1",
            "Address:        192.168.10.1#53",
            "",
            "www.pasundan.sch.id canonical name = pasundan.sch.id.",
            "Name:   pasundan.sch.id",
            "Address: 192.168.10.1"
        ],
        "a": "A. Berhasil meresolusi domain dan CNAME ke IP 192.168.10.1 pada port 53",
        "b": "B. Gagal karena server menolak query (REFUSED)",
        "c": "C. Timeout karena firewall memblokir port 80",
        "d": "D. Domain tidak terdaftar di root server",
        "e": "E. Terjadi loop query pada DNS relay",
        "kunci": "A"
    },
    {
        "cat": "DNS SERVER",
        "title": "Konfigurasi Klien DNS (/etc/resolv.conf)",
        "q": "File konfigurasi pada klien Linux yang bertugas menentukan alamat IP DNS resolver lokal yang dituju adalah...",
        "lines": [
            "root@debian12:~# cat /etc/resolv.conf",
            "nameserver 192.168.10.1",
            "nameserver 8.8.8.8",
            "search pasundan.sch.id"
        ],
        "a": "A. /etc/resolv.conf",
        "b": "B. /etc/hosts.allow",
        "c": "C. /etc/hostname",
        "d": "D. /etc/network/interfaces",
        "e": "E. /etc/fstab",
        "kunci": "A"
    },
    {
        "cat": "DNS SERVER",
        "title": "Sintaks Validasi Konfigurasi BIND9",
        "q": "Perintah bawaan BIND9 untuk memeriksa apakah ada kesalahan sintaks penulisan pada file named.conf sebelum me-restart service adalah...",
        "lines": [
            "root@debian12:~# named-checkconf",
            "root@debian12:~# named-checkzone pasundan.sch.id /etc/bind/db.pasundan",
            "zone pasundan.sch.id/IN: loaded serial 2026082901",
            "OK"
        ],
        "a": "A. named-checkconf dan named-checkzone",
        "b": "B. bind9-test --verify",
        "c": "C. apache2ctl configtest",
        "d": "D. postfix check",
        "e": "E. dhcpd -t",
        "kunci": "A"
    },
    {
        "cat": "DNS SERVER",
        "title": "Port Default Protokol DNS",
        "q": "DNS Server beroperasi secara default pada port jaringan dan protokol...",
        "lines": [
            "Protocol: UDP / TCP",
            "Default Port: 53",
            "Service Daemon: named (BIND9)"
        ],
        "a": "A. Port 53 (UDP & TCP)",
        "b": "B. Port 80 (TCP)",
        "c": "C. Port 443 (TCP)",
        "d": "D. Port 21 (TCP)",
        "e": "E. Port 25 (TCP)",
        "kunci": "A"
    },

    # 15 - 22: Web Server (Apache2)
    {
        "cat": "WEB SERVER (APACHE2)",
        "title": "Virtual Host Apache2 Configuration",
        "q": "Perhatikan konfigurasi Virtual Host Apache2 berikut di /etc/apache2/sites-available/pasundan.conf:\nDirektori penyimpanan file index.html / website yang akan dimuat saat domain diakses adalah...",
        "lines": [
            "<VirtualHost *:80>",
            "    ServerAdmin admin@pasundan.sch.id",
            "    ServerName pasundan.sch.id",
            "    ServerAlias www.pasundan.sch.id",
            "    DocumentRoot /var/www/pasundan",
            "    ErrorLog ${APACHE_LOG_DIR}/pasundan_error.log",
            "    CustomLog ${APACHE_LOG_DIR}/pasundan_access.log combined",
            "</VirtualHost>"
        ],
        "a": "A. /var/www/pasundan",
        "b": "B. /etc/apache2/sites-available",
        "c": "C. /var/log/apache2",
        "d": "D. /usr/share/apache2",
        "e": "E. /etc/apache2/mods-enabled",
        "kunci": "A"
    },
    {
        "cat": "WEB SERVER (APACHE2)",
        "title": "Mengaktifkan Virtual Host (a2ensite)",
        "q": "Perintah CLI untuk mengaktifkan konfigurasi Virtual Host pasundan.conf pada Apache2 di Debian 12 adalah...",
        "lines": [
            "root@debian12:~# a2ensite pasundan.conf",
            "Enabling site pasundan.",
            "To activate the new configuration, you need to run:",
            "  systemctl reload apache2"
        ],
        "a": "A. a2ensite pasundan.conf",
        "b": "B. a2dissite pasundan.conf",
        "c": "C. a2enmod rewrite",
        "d": "D. apache2ctl start site",
        "e": "E. apt install pasundan.conf",
        "kunci": "A"
    },
    {
        "cat": "WEB SERVER (APACHE2)",
        "title": "Mengaktifkan Modul Rewrite & SSL (a2enmod)",
        "q": "Untuk mengaktifkan modul SSL (HTTPS) dan Rewrite URL pada Apache2 di Debian 12, perintah yang digunakan adalah...",
        "lines": [
            "root@debian12:~# a2enmod ssl",
            "root@debian12:~# a2enmod rewrite",
            "root@debian12:~# a2ensite default-ssl.conf",
            "root@debian12:~# systemctl restart apache2"
        ],
        "a": "A. a2enmod ssl && a2enmod rewrite",
        "b": "B. apt install apache2-ssl-rewrite",
        "c": "C. a2ensite rewrite.conf",
        "d": "D. systemctl enable ssl",
        "e": "E. chmod +x ssl rewrite",
        "kunci": "A"
    },
    {
        "cat": "WEB SERVER",
        "title": "Port Standar HTTP & HTTPS",
        "q": "Port default protokol web standar (HTTP) dan web terenkripsi (HTTPS) secara berurutan adalah...",
        "lines": [
            "HTTP  (Unencrypted Web): Port 80 / TCP",
            "HTTPS (SSL/TLS Encrypted): Port 443 / TCP",
            "Configuration file: /etc/apache2/ports.conf"
        ],
        "a": "A. Port 80 (HTTP) dan Port 443 (HTTPS)",
        "b": "B. Port 8080 (HTTP) dan Port 8443 (HTTPS)",
        "c": "C. Port 21 (HTTP) dan Port 22 (HTTPS)",
        "d": "D. Port 25 (HTTP) dan Port 110 (HTTPS)",
        "e": "E. Port 53 (HTTP) dan Port 67 (HTTPS)",
        "kunci": "A"
    },
    {
        "cat": "WEB SERVER (APACHE2)",
        "title": "File Konfigurasi Utama Apache2",
        "q": "File konfigurasi global utama server Apache2 pada sistem operasi Debian 12 terletak di...",
        "lines": [
            "root@debian12:~# ls -la /etc/apache2/",
            "apache2.conf       # Global configuration",
            "ports.conf         # Listen ports (80, 443)",
            "sites-available/   # VirtualHost configs",
            "sites-enabled/     # Symlinks to active sites",
            "mods-available/    # Available modules"
        ],
        "a": "A. /etc/apache2/apache2.conf",
        "b": "B. /etc/httpd/conf/httpd.conf",
        "c": "C. /var/www/html/apache2.conf",
        "d": "D. /usr/sbin/apache2.conf",
        "e": "E. /etc/nginx/nginx.conf",
        "kunci": "A"
    },
    {
        "cat": "WEB SERVER (APACHE2)",
        "title": "Uji Konfigurasi Sintaks Apache (apache2ctl configtest)",
        "q": "Perintah untuk memverifikasi apakah ada kesalahan sintaks pada seluruh file konfigurasi Apache2 sebelum direload adalah...",
        "lines": [
            "root@debian12:~# apache2ctl configtest",
            "AH00558: apache2: Could not reliably determine the server's fully qualified domain name",
            "Syntax OK"
        ],
        "a": "A. apache2ctl configtest",
        "b": "B. apache2 --version",
        "c": "C. a2check site",
        "d": "D. service apache2 test",
        "e": "E. systemctl check apache2",
        "kunci": "A"
    },
    {
        "cat": "WEB SERVER",
        "title": "Directory Index (Urutan Pemuatan File Web)",
        "q": "Pada modul dir.conf Apache2, file yang akan dicari dan dieksekusi pertama kali oleh web server saat klien mengakses alamat web tanpa menyebutkan nama file adalah...",
        "lines": [
            "<IfModule mod_dir.c>",
            "    DirectoryIndex index.php index.html index.cgi index.pl",
            "</IfModule>"
        ],
        "a": "A. index.php (jika ada), lalu index.html",
        "b": "B. default.asp",
        "c": "C. home.html",
        "d": "D. index.cgi selalu yang pertama",
        "e": "E. error.html",
        "kunci": "A"
    },
    {
        "cat": "WEB SERVER",
        "title": "Log Akses & Error Apache2",
        "q": "Untuk memantau log error web server Apache2 secara real-time saat terjadi troubleshooting, perintah CLI yang tepat adalah...",
        "lines": [
            "root@debian12:~# tail -f /var/log/apache2/error.log",
            "[Sat Aug 29 07:15:20 2026] [mpm_event:notice] [pid 1234:tid 5678] AH00489: Apache/2.4.59 (Debian) configured -- resuming normal operations"
        ],
        "a": "A. tail -f /var/log/apache2/error.log",
        "b": "B. head /var/log/apache2/error.log",
        "c": "C. nano /etc/apache2/error.log",
        "d": "D. rm /var/log/apache2/error.log",
        "e": "E. cat /var/log/syslog | grep dhcp",
        "kunci": "A"
    },

    # 23 - 28: DHCP Server (ISC-DHCP-SERVER)
    {
        "cat": "DHCP SERVER",
        "title": "Menentukan Interface DHCP (/etc/default/isc-dhcp-server)",
        "q": "Pada Debian 12, file konfigurasi yang digunakan untuk menentukan interface kartu jaringan mana yang akan menyiarkan layanan DHCP (misal: ens18) adalah...",
        "lines": [
            "root@debian12:~# cat /etc/default/isc-dhcp-server",
            "# On what interfaces should the DHCP server serve DHCP requests?",
            "INTERFACESv4=\"ens18\"",
            "INTERFACESv6=\"\""
        ],
        "a": "A. /etc/default/isc-dhcp-server",
        "b": "B. /etc/dhcp/dhcpd.conf",
        "c": "C. /etc/network/interfaces",
        "d": "D. /etc/resolv.conf",
        "e": "E. /var/lib/dhcp/dhcpd.leases",
        "kunci": "A"
    },
    {
        "cat": "DHCP SERVER",
        "title": "Konfigurasi Subnet & Range IP (/etc/dhcp/dhcpd.conf)",
        "q": "Perhatikan konfigurasi file /etc/dhcp/dhcpd.conf berikut:\nRentang IP address yang akan dibagikan secara otomatis kepada klien adalah...",
        "lines": [
            "subnet 192.168.10.0 netmask 255.255.255.0 {",
            "    range 192.168.10.50 192.168.10.100;",
            "    option routers 192.168.10.1;",
            "    option domain-name-servers 192.168.10.1, 8.8.8.8;",
            "    default-lease-time 600;",
            "    max-lease-time 7200;",
            "}"
        ],
        "a": "A. 192.168.10.50 sampai 192.168.10.100 (51 Host)",
        "b": "B. 192.168.10.1 sampai 192.168.10.254",
        "c": "C. Hanya IP 192.168.10.1",
        "d": "D. 192.168.10.0 sampai 192.168.10.255",
        "e": "E. 8.8.8.8 sampai 8.8.4.4",
        "kunci": "A"
    },
    {
        "cat": "DHCP SERVER",
        "title": "DHCP Option Routers (Default Gateway)",
        "q": "Parameter 'option routers' pada file konfigurasi dhcpd.conf berfungsi untuk memberitahukan kepada klien alamat...",
        "lines": [
            "option routers 192.168.10.1;          # Default Gateway",
            "option domain-name \"pasundan.sch.id\"; # Domain Search",
            "option domain-name-servers 192.168.10.1; # DNS Server"
        ],
        "a": "A. Default Gateway (Pintu Keluar Jaringan)",
        "b": "B. Web Server",
        "c": "C. Database Server",
        "d": "D. Mail Server",
        "e": "E. Subnet Mask",
        "kunci": "A"
    },
    {
        "cat": "DHCP SERVER",
        "title": "DHCP Reservation (Static Leases / MAC Address Binding)",
        "q": "Administrator ingin agar komputer Kepala Sekolah dengan MAC address 08:00:27:12:34:56 selalu mendapatkan IP khusus 192.168.10.10. Konfigurasi yang tepat pada dhcpd.conf adalah...",
        "lines": [
            "host PC-Kepsek {",
            "    hardware ethernet 08:00:27:12:34:56;",
            "    fixed-address 192.168.10.10;",
            "}"
        ],
        "a": "A. hardware ethernet 08:00:27:12:34:56; fixed-address 192.168.10.10;",
        "b": "B. range 192.168.10.10 192.168.10.10;",
        "c": "C. option static-ip 192.168.10.10;",
        "d": "D. mac-address 08:00:27:12:34:56 -> 192.168.10.10;",
        "e": "E. default-lease-time fixed;",
        "kunci": "A"
    },
    {
        "cat": "DHCP DORA PROCESS",
        "title": "Urutan Proses Alokasi IP DHCP (D-O-R-A)",
        "q": "Urutan 4 tahap komunikasi paket antara DHCP Client dan DHCP Server saat meminta IP address adalah...",
        "lines": [
            "1. Client -> Server : DHCP DISCOVER (Broadcast mencari server)",
            "2. Server -> Client : DHCP OFFER (Penawaran IP)",
            "3. Client -> Server : DHCP REQUEST (Meminta IP yang ditawarkan)",
            "4. Server -> Client : DHCP ACKNOWLEDGEMENT (Konfirmasi & Sewa IP)"
        ],
        "a": "A. Discover, Offer, Request, Acknowledge (DORA)",
        "b": "B. Request, Discover, Offer, Acknowledge",
        "c": "C. Offer, Request, Discover, Acknowledge",
        "d": "D. Acknowledge, Request, Offer, Discover",
        "e": "E. Discover, Request, Offer, Acknowledge",
        "kunci": "A"
    },
    {
        "cat": "DHCP SERVER",
        "title": "Port Default Layanan DHCP",
        "q": "DHCP Server dan DHCP Client berkomunikasi menggunakan protokol UDP pada port...",
        "lines": [
            "DHCP Server : Port 67 / UDP",
            "DHCP Client : Port 68 / UDP"
        ],
        "a": "A. Port 67 (Server) dan Port 68 (Client)",
        "b": "B. Port 80 (Server) dan Port 8080 (Client)",
        "c": "C. Port 20 (Server) dan Port 21 (Client)",
        "d": "D. Port 25 (Server) dan Port 110 (Client)",
        "e": "E. Port 53 (Server) dan Port 53 (Client)",
        "kunci": "A"
    },

    # 29 - 34: FTP Server (ProFTPD / VSFTPD)
    {
        "cat": "FTP SERVER (VSFTPD)",
        "title": "Konfigurasi Keamanan Chroot (/etc/vsftpd.conf)",
        "q": "Pada konfigurasi /etc/vsftpd.conf, opsi yang harus diaktifkan agar user FTP yang login terkunci di dalam direktori home-nya masing-masing dan tidak dapat menjelajah ke direktori root sistem (/etc, /var, dll) adalah...",
        "lines": [
            "root@debian12:~# nano /etc/vsftpd.conf",
            "anonymous_enable=NO",
            "local_enable=YES",
            "write_enable=YES",
            "chroot_local_user=YES",
            "allow_writeable_chroot=YES"
        ],
        "a": "A. chroot_local_user=YES",
        "b": "B. anonymous_enable=YES",
        "c": "C. write_enable=NO",
        "d": "D. local_enable=NO",
        "e": "E. listen_port=2121",
        "kunci": "A"
    },
    {
        "cat": "FTP SERVER",
        "title": "Port Standar Protokol FTP (Control & Data)",
        "q": "Protokol FTP beroperasi menggunakan dua port TCP utama, yaitu port untuk jalur perintah (command/control) dan port untuk transfer data. Kedua port tersebut adalah...",
        "lines": [
            "FTP Control (Command channel) : Port 21 / TCP",
            "FTP Data (Data transfer)       : Port 20 / TCP"
        ],
        "a": "A. Port 21 (Control) dan Port 20 (Data)",
        "b": "B. Port 22 (Control) dan Port 23 (Data)",
        "c": "C. Port 80 (Control) dan Port 443 (Data)",
        "d": "D. Port 25 (Control) dan Port 110 (Data)",
        "e": "E. Port 53 (Control) dan Port 67 (Data)",
        "kunci": "A"
    },
    {
        "cat": "FTP SERVER",
        "title": "Menonaktifkan Akses Anonymous FTP",
        "q": "Untuk mencegah pengguna asing mengunduh atau mengunggah file tanpa password (user anonim), konfigurasi yang harus diset pada vsftpd.conf adalah...",
        "lines": [
            "# Disable unauthenticated guest access",
            "anonymous_enable=NO",
            "local_enable=YES"
        ],
        "a": "A. anonymous_enable=NO",
        "b": "B. anonymous_enable=YES",
        "c": "C. guest_enable=YES",
        "d": "D. no_anon_password=YES",
        "e": "E. allow_anon_write=YES",
        "kunci": "A"
    },
    {
        "cat": "FTP SERVER",
        "title": "Mengizinkan Upload File pada FTP",
        "q": "Jika user berhasil login ke FTP Server tetapi mendapatkan error '550 Permission Denied' saat mencoba mengunggah file, opsi konfigurasi yang lupa diaktifkan pada vsftpd.conf adalah...",
        "lines": [
            "root@debian12:~# nano /etc/vsftpd.conf",
            "# Uncomment this to enable any form of FTP write command.",
            "write_enable=YES"
        ],
        "a": "A. write_enable=YES",
        "b": "B. read_only=YES",
        "c": "C. upload_mode=ON",
        "d": "D. file_access=WRITE",
        "e": "E. chown_uploads=NO",
        "kunci": "A"
    },
    {
        "cat": "FTP SERVER",
        "title": "Perbedaan FTP vs SFTP",
        "q": "Perbedaan mendasar antara protokol FTP murni (vsftpd/proftpd) dan SFTP (SSH File Transfer Protocol) dalam aspek keamanan adalah...",
        "lines": [
            "FTP  : Plaintext (Username, password, dan data tidak terenkripsi, Port 21)",
            "SFTP : Encrypted (Berjalan di atas protokol SSH terenkripsi, Port 22)"
        ],
        "a": "A. SFTP mengenkripsi seluruh transfer data dan password melalui SSH (Port 22)",
        "b": "B. FTP lebih aman daripada SFTP",
        "c": "C. SFTP hanya bisa digunakan di Windows",
        "d": "D. FTP menggunakan enkripsi AES 256 secara default",
        "e": "E. SFTP tidak memerlukan user dan password",
        "kunci": "A"
    },
    {
        "cat": "FTP CLIENT",
        "title": "Aplikasi Klien FTP Populer",
        "q": "Software aplikasi klien GUI populer yang sering digunakan di Windows untuk mentransfer file ke FTP Server Debian 12 adalah...",
        "lines": [
            "FileZilla Client / WinSCP",
            "Host: 192.168.10.1, Username: siswa, Port: 21",
            "Status: Directory listing of \"/\" successful"
        ],
        "a": "A. FileZilla dan WinSCP",
        "b": "B. Wireshark",
        "c": "C. VirtualBox",
        "d": "D. Cisco Packet Tracer",
        "e": "E. PuTTY terminal",
        "kunci": "A"
    },

    # 35 - 40: Mail Server (Postfix & Dovecot)
    {
        "cat": "MAIL SERVER (POSTFIX)",
        "title": "Konfigurasi Domain Mail (/etc/postfix/main.cf)",
        "q": "Perhatikan konfigurasi Postfix pada file /etc/postfix/main.cf berikut:\nParameter yang menentukan nama domain resmi yang ditangani oleh server email ini adalah...",
        "lines": [
            "myhostname = mail.pasundan.sch.id",
            "mydomain = pasundan.sch.id",
            "myorigin = $mydomain",
            "mydestination = $myhostname, pasundan.sch.id, localhost.localdomain, localhost",
            "home_mailbox = Maildir/"
        ],
        "a": "A. mydomain = pasundan.sch.id dan mydestination",
        "b": "B. inet_interfaces = loopback-only",
        "c": "C. smtpd_banner = $myhostname ESMTP",
        "d": "D. alias_maps = hash:/etc/aliases",
        "e": "E. mailbox_size_limit = 0",
        "kunci": "A"
    },
    {
        "cat": "MAIL SERVER (POSTFIX)",
        "title": "Format Penyimpanan Email (Maildir vs Mbox)",
        "q": "Konfigurasi 'home_mailbox = Maildir/' pada Postfix memiliki keunggulan dibandingkan format mbox, yaitu...",
        "lines": [
            "Format Mbox    : Semua email ditumpuk menjadi 1 file tunggal besar (/var/mail/user)",
            "Format Maildir : Setiap email disimpan sebagai file terpisah di direktori ~/Maildir/ (cur, new, tmp)"
        ],
        "a": "A. Setiap pesan email disimpan sebagai file tersendiri di folder user sehingga lebih aman dari korupsi data",
        "b": "B. Menghapus seluruh email secara otomatis setelah 7 hari",
        "c": "C. Mengompresi file email menjadi file ZIP",
        "d": "D. Mengubah protokol SMTP menjadi HTTP",
        "e": "E. Tidak membutuhkan kapasitas harddisk",
        "kunci": "A"
    },
    {
        "cat": "MAIL SERVER (DOVECOT)",
        "title": "Protokol Pengambilan Email (IMAP vs POP3)",
        "q": "Dovecot menyediakan dua protokol utama bagi klien untuk mengambil email: IMAP dan POP3. Karakteristik utama dari protokol IMAP adalah...",
        "lines": [
            "IMAP (Port 143 / SSL 993) : Email tetap tersimpan di server dan tersinkronisasi di semua perangkat",
            "POP3 (Port 110 / SSL 995) : Email diunduh ke perangkat klien dan dihapus dari server"
        ],
        "a": "A. Email tersinkronisasi dua arah dan tetap tersimpan di server",
        "b": "B. Email langsung dihapus permanen dari server setelah diunduh",
        "c": "C. Hanya bisa mengirim email tanpa bisa menerima",
        "d": "D. Tidak mendukung koneksi terenkripsi SSL",
        "e": "E. Menggunakan port 25",
        "kunci": "A"
    },
    {
        "cat": "MAIL SERVER",
        "title": "Port Standar SMTP, POP3, dan IMAP",
        "q": "Port default standar tanpa enkripsi untuk protokol pengiriman email (SMTP), pengambilan email (POP3), dan sinkronisasi email (IMAP) secara berurutan adalah...",
        "lines": [
            "SMTP (Kirim Email)  : Port 25 / TCP",
            "POP3 (Tarik Email)  : Port 110 / TCP",
            "IMAP (Akses Email)  : Port 143 / TCP"
        ],
        "a": "A. SMTP: 25, POP3: 110, IMAP: 143",
        "b": "B. SMTP: 80, POP3: 443, IMAP: 21",
        "c": "C. SMTP: 53, POP3: 67, IMAP: 68",
        "d": "D. SMTP: 22, POP3: 23, IMAP: 25",
        "e": "E. SMTP: 465, POP3: 995, IMAP: 993",
        "kunci": "A"
    },
    {
        "cat": "WEBMAIL",
        "title": "Aplikasi Webmail (Roundcube / Squirrelmail)",
        "q": "Software antarmuka berbasis web (Webmail Client) yang sering diintegrasikan dengan Apache2, Postfix, dan Dovecot agar pengguna dapat mengirim dan membaca email melalui browser adalah...",
        "lines": [
            "root@debian12:~# apt install roundcube roundcube-mysql",
            "Browser URL: http://mail.pasundan.sch.id/roundcube",
            "Webmail Engine: Roundcube Webmail / Squirrelmail"
        ],
        "a": "A. Roundcube Webmail dan Squirrelmail",
        "b": "B. Wireshark",
        "c": "C. phpMyAdmin",
        "d": "D. Cisco IOS",
        "e": "E. VirtualBox Manager",
        "kunci": "A"
    },
    {
        "cat": "MAIL SERVER",
        "title": "Peran MX Record pada DNS untuk Mail Server",
        "q": "Mengapa sebuah Mail Server (Postfix) memerlukan konfigurasi MX (Mail Exchanger) Record pada DNS Server?",
        "lines": [
            "; DNS Zone pasundan.sch.id",
            "@       IN      MX  10  mail.pasundan.sch.id.",
            "mail    IN      A       192.168.10.1"
        ],
        "a": "A. Agar server email pengirim dari luar internet mengetahui host mana yang bertanggung jawab menerima email untuk domain tersebut",
        "b": "B. Untuk membatasi ukuran file lampiran (attachment) email",
        "c": "C. Untuk mengubah password user email secara otomatis",
        "d": "D. Untuk menghubungkan virtual host Apache ke port 80",
        "e": "E. Untuk mengalokasikan IP address DHCP kepada komputer klien",
        "kunci": "A"
    }
]

def generate_linux_docx():
    print("🚀 Generating 40 Linux, DNS, Web, Apache2, DHCP, FTP, Mail questions on Debian 12 & VirtualBox...")
    doc = Document()

    # Title & Header
    title = doc.add_heading("BANK SOAL TEKNIK KOMPUTER & JARINGAN (40 BUTIR SOAL)", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Topik: LINUX DEBIAN 12, VIRTUALBOX, DNS (BIND9), WEB (APACHE2), DHCP, FTP, & MAIL SERVER")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub2 = doc.add_paragraph("Disertai 40 Diagram Terminal CLI, Skema Jaringan & Konfigurasi Beresolusi Tinggi")
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("=" * 80)

    for i, item in enumerate(QUESTIONS):
        img_name = f"linux_q{i+1}.png"
        img_path = create_terminal_image(img_name, item["cat"], item["title"], item["lines"])

        p = doc.add_paragraph()
        p.add_run(f"{i+1}. [{item['cat']}] {item['q']}").bold = True

        doc.add_picture(img_path, width=Inches(5.0))

        doc.add_paragraph(item["a"])
        doc.add_paragraph(item["b"])
        doc.add_paragraph(item["c"])
        doc.add_paragraph(item["d"])
        doc.add_paragraph(item["e"])
        
        p_key = doc.add_paragraph()
        p_key.add_run(f"KUNCI: {item['kunci']}").bold = True
        doc.add_paragraph()

    # Save to all target paths
    out_scratch = os.path.join(os.path.dirname(__file__), "../BANK_SOAL_40_LINUX_DEBIAN12_VBOX.docx")
    out_public = os.path.join(os.path.dirname(__file__), "../public/BANK_SOAL_40_LINUX_DEBIAN12_VBOX.docx")
    out_desktop = os.path.expanduser("~/Desktop/BANK_SOAL_40_LINUX_DEBIAN12_VBOX.docx")
    out_downloads = os.path.expanduser("~/Downloads/BANK_SOAL_40_LINUX_DEBIAN12_VBOX.docx")

    doc.save(out_scratch)
    shutil.copy(out_scratch, out_public)
    try:
        shutil.copy(out_scratch, out_desktop)
        shutil.copy(out_scratch, out_downloads)
    except Exception as e:
        print(f"Copy warning: {e}")

    print(f"✅ Dokumen 40 Soal Linux Debian 12 & VirtualBox berhasil dibuat di:\n - {out_desktop}\n - {out_downloads}\n - {out_scratch}")

if __name__ == "__main__":
    generate_linux_docx()
